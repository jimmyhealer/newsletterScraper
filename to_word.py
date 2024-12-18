import docx
from docx.shared import Pt, Cm
import pathlib
from rich.progress import Progress
from utils import check_filename_legal
import concurrent.futures
import logging
import sys
import re

def get_executable_path():
    if getattr(sys, 'frozen', False):
        executable_path = pathlib.Path(sys.executable).resolve().parent
    else:
        executable_path = pathlib.Path(__file__).resolve().parent
    return executable_path

def data_to_word(datas):
    script_dir = get_executable_path()
    output_dir = script_dir / "output"
    output_dir.mkdir(exist_ok=True)

    with Progress() as progress:
        task = progress.add_task("[cyan] 正在寫入 Word 文件...", total=len(datas))

        # 使用 ThreadPoolExecutor 進行多執行緒處理
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = []

            for data in datas:
                # 將任務提交到執行緒池
                future = executor.submit(write_to_word_task, output_dir, data)
                futures.append(future)

            # 更新進度條，當任務完成時
            for future in concurrent.futures.as_completed(futures):
                progress.update(task, advance=1)


def write_to_word_task(root, data):
    # 建立一個 Word 文件
    doc = docx.Document()

    # 標題 - 粗體 16pt，標楷體
    title = doc.add_paragraph()
    title_run = title.add_run(data["title"])
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "標楷體"

    # 設置中文字體（解決中文字體設置問題）
    rFonts = title_run._element.rPr.rFonts
    rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "標楷體")

    # 標籤 - 12pt，標楷體
    labels_paragraph = doc.add_paragraph()
    labels_run = labels_paragraph.add_run(f"{data['label']} ／ {data['author']}")
    labels_run.font.size = Pt(12)
    labels_run.font.name = "標楷體"

    # 設置中文字體
    rFonts = labels_run._element.rPr.rFonts
    rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "標楷體")

    # 插入圖片
    for images in data["images"]:
        try:
            doc.add_picture(images, width=Cm(16))
        except Exception as e:
            logging.error(f"Error inserting image {images}: {e}")

    # 內文 - 12pt，標楷體
    content_paragraph = doc.add_paragraph()
    content_parts = re.split(r"(\*\*.*?\*\*)", data["content"])

    for part in content_parts:
        content_run = content_paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        content_run.font.size = Pt(12)
        content_run.font.name = "標楷體"
        if part.startswith("**") and part.endswith("**"):
            content_run.bold = True

        # 設置中文字體
        rFonts = content_run._element.rPr.rFonts
        rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "標楷體")

    # 儲存 Word 文件
    filename = root / check_filename_legal(f"電子時報_{data['date']}") / check_filename_legal(f"{data['date']}_{data['title']}.docx")
    filename.parent.mkdir(exist_ok=True)
    doc.save(filename)
